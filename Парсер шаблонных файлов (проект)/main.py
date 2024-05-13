import os
import shutil
import sqlite3
import tkinter as tk
from tkinter import ttk
import re
import docx
import win32clipboard
import json
import csv
import pandas as pd
import sys


def to_json(data, filename):
    with open(filename, "w") as f:
        json.dump(data, f)


def to_csv(data, filename):
    with open(filename, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(data.keys())
        writer.writerow(data.values())


def to_excel(data, filename):
    df = pd.DataFrame([data])
    writer = pd.ExcelWriter(filename)
    df.to_excel(writer, index=False)
    writer.save()


def find_docx_shablons():
    """
    Функция возвращает список файлов-шаблонов в папке проекта/templates
    """
    templates_path = os.path.join(os.getcwd(), 'templates')

    files = []
    for file in os.listdir(templates_path):
        if 'shablon' in file:
            files.append(file)
    return files


def read_paragraphs(doc_path):
    """
    Функция возвращает список параграфов текста (note: игнорирует таблицы)
    :param doc_path: путь к анализиуемому файлу

    :return: список строк, где каждая строка - один параграф
    """

    paragraphs = []
    document = docx.Document(doc_path)
    for i in document.paragraphs:
        paragraphs.append(i.text)
    return paragraphs


def replace_variables(string: str, separator_left: str, separator_right: str):
    """
    Функция заменяет сочетания символов, окруженные сепараторами (метки), на символ & (например, "основанный в <YEAR> году" -> "основанный в & году")
    :param string: анализируемая строка
    :param separator_left: сепаратор слева от заменяемого сочетания символов
    :param separator_right: сепаратор справа от заменяемого сочетания символов

    :return: входная строка с замененными метками на символ &
    """
    return re.sub(rf'{separator_left}.*?{separator_right}', '&', string)


def get_data_from_string(template: str, string: str, separator_left: str, separator_right: str):
    """
    Функция извлекает данные из шаблонной строки методом вычитания статичных сочетаний символов:

    :param template: размеченная строка-шаблон (например, "<SELLER>, выступающий от лица компании <COMPANY_NAME> на основании <ORDER_NUMBER>", где <SELLER>, <COMPANY_NAME> и <ORDER_NUMBER> стоят на месте изменяемых частей строки, а ", выступающий от лица компании " и " на основании " - статичные части строки)
    :param string: анализируемая строка документа, из которой необходимо извлечь данные по шаблону. Например, "Иванов Иван Иванович, выступающий от лица компании 3D PlastPrint на основании приказа №IN027/SEL0523", где фрагменты "Иванов Иван Иванович", "3D PlastPrint" и "приказа №IN027/SEL0523" будут извлечены под именами переменных <SELLER>, <COMPANY_NAME> и <ORDER_NUMBER> соответственно.
    :param separator_left: символ или сочетание символов, ограничивающее метку переменной слева. В приведенном выше примере это "<"
    :param separator_right: символ или сочетание символов, ограничивающее метку переменной справа. В приведенном выше примере это ">"

    :return: словарь вида {VARIABLE_NAME: VARIABLE_VALUE}, где VARIABLE_NAME - имя переменной, полученное из размеченного шаблона, VARIABLE_VALUE - значение данной переменной, полученное из анализируемой строки. В случае приведенного выше примера будет возвращен словарь {"SELLER": "Иванов Иван Иванович", "COMPANY_NAME": "3D PlastPrint", "ORDER_NUMBER": "приказа №IN027/SEL0523"}
    """

    def find_tags_in_string(string, sep_l, sep_r):
        """
        Функция возвращает список всех сочетаний символов, окруженных сепараторами (меток), (например, "основанный в <YEAR> году" -> ["YEAR"]):

        :param string: анализируемая строка
        :param sep_left: сепаратор слева от метки
        :param sep_right: сепаратор справа от метки

        :return: список тегов в строке
        """

        tag_pattern = re.compile(rf'{sep_l}.*?{sep_r}')

        tags = tag_pattern.findall(string)

        result_tags = []
        for tag in tags:
            result_tags.append(tag.replace(sep_l, '').replace(sep_r, ''))

        return result_tags

    static = replace_variables(string=template, separator_left=separator_left, separator_right=separator_right)
    for i in static.split('&'):
        if i != '' and i in string:
            string = string.replace(i, '|')
    values = string.split('|')
    variables = find_tags_in_string(string=template, sep_l=separator_left, sep_r=separator_right)
    dict = {}
    for elem in variables:
        if elem == '':
            variables.remove(elem)
    for elem in values:
        if elem == '':
            values.remove(elem)

    if len(values) == len(variables):
        i = 0
        while i < len(values):
            dict[variables[i]] = values[i]
            i += 1
    return dict


def find_tags_in_string(string, sep_l, sep_r):
    """
    Функция возвращает список всех сочетаний символов, окруженных сепараторами (меток), (например, "основанный в <YEAR> году" -> ["YEAR"])
    :param string: анализируемая строка
    :param sep_left: сепаратор слева от метки
    :param sep_right: сепаратор справа от метки

    :return: список тегов в строке
    """

    tag_pattern = re.compile(rf'{sep_l}.*?{sep_r}')

    tags = tag_pattern.findall(string)

    result_tags = []
    for tag in tags:
        result_tags.append(tag.replace(sep_l, '').replace(sep_r, ''))

    return result_tags


def extract_tables_from_docx(file_path):
    """
    Функция извлекает содержимое ячеек всех таблиц файла Word
    :param file_path: путь к файлу, содержимое таблиц которого необходимо извлечь

    :return: словарь вида {'порядковый номер таблицы в документе': {'номер строки*номер столбца': содержимое ячейки}}
    """

    doc = docx.Document(file_path)
    tables_data = {}

    table_index = 1
    for table in doc.tables:
        table_data = {}
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                table_data[f'{row_index + 1}*{cell_index + 1}'] = cell.text
        tables_data[table_index] = table_data
        table_index += 1

    return tables_data


class App:
    def __init__(self):
        if os.path.exists("settings.db"):
            try:
                self.conn = sqlite3.connect("settings.db")
                self.cursor = self.conn.cursor()
            except:
                print('Не удалось открыть базу данных')
        else:
            try:
                self.conn = sqlite3.connect("settings.db")
                self.cursor = self.conn.cursor()
                self.cursor.execute('''CREATE TABLE templates (template text, separator_left, separator_right text)''')
                self.conn.commit()
            except:
                print("Не удалось создать базу данных")

        self.root = tk.Tk()
        self.root.title("v 1.0.2")
        self.root.geometry("460x300")
        self.root.resizable(width=False, height=False)

        self.label = tk.Label(self.root, text="Программа готова к работе.\n Выберите необходимое действие в меню")
        self.label.place(relx=0.5, rely=0.5, anchor="center")
        self.main_menu()
        self.root.mainloop()

    def templates(self):
        def find_template(file, sep_l, sep_r):
            def copy_to_templates(source_file_path):
                try:
                    project_dir = os.path.dirname(os.path.abspath(__file__))
                    templates_dir = os.path.join(project_dir, "templates")

                    os.makedirs(templates_dir, exist_ok=True)

                    source_file_name = os.path.basename(source_file_path)

                    dest_file_path = os.path.join(templates_dir, source_file_name)

                    shutil.copyfile(source_file_path, dest_file_path)

                    return os.path.relpath(dest_file_path, start=project_dir).replace("\\", "/")
                except FileNotFoundError:
                    print("Файл не найден:", source_file_path)
                except PermissionError:
                    print("Ошибка доступа к файлу:", source_file_path)
                except shutil.SameFileError:
                    print("Нельзя скопировать файл сам в себя:", source_file_path)
                except shutil.Error as e:
                    print("Ошибка при копировании файла:", e)
                except Exception as e:
                    print("Неизвестная ошибка:", e)

            copied_file = copy_to_templates(file)

            paragraph_data = []
            table_data = []

            if copied_file.startswith('templates'):
                conn = sqlite3.connect('settings.db')
                c = conn.cursor()

                c.execute(f"DELETE FROM templates WHERE template LIKE '%{copied_file}%'")
                c.execute("INSERT INTO templates VALUES(?, ?, ?)", (copied_file, sep_l, sep_r))

                conn.commit()
                conn.close()

                tables_data = extract_tables_from_docx(file_path=file)
                tables = list(tables_data.keys())
                for i in tables:
                    for a in list(tables_data[i].keys()):
                        cell_static_text = replace_variables(string=tables_data[i][a], separator_left=sep_l,
                                                             separator_right=sep_r).split("&")
                        for b in cell_static_text:
                            table_data.append((copied_file, i, a, b))

                paragraphs = read_paragraphs(doc_path=file)
                for paragraph in paragraphs:
                    if paragraph != '':
                        static_variables = replace_variables(string=paragraph, separator_left=sep_l,
                                                             separator_right=sep_r).split('&')
                        for i in static_variables:
                            paragraph_data.append((copied_file, paragraph, i))
            if paragraph_data != []:
                table = ttk.Treeview(self.root, columns=("Файл", "Параграф", "Статичный текст"), show="headings")

                for i in paragraph_data:
                    table.insert("", tk.END, values=i)

                table.column("Файл", width=140, anchor=tk.CENTER)
                table.column("Параграф", width=140, anchor=tk.CENTER)
                table.column("Статичный текст", width=140, anchor=tk.CENTER)

                table.heading("Файл", text="Файл")
                table.heading("Параграф", text="Параграф")
                table.heading("Статичный текст", text="Статичный текст")

                scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=table.yview)

                table.configure(yscrollcommand=scrollbar.set)

                table.place(x=10, y=160, height=110)
                scrollbar.place(x=435, y=160, height=110)
            elif table_data != []:
                table = ttk.Treeview(self.root, columns=("Файл", "Порядковый номер таблицы",
                                                         "Номер строки*Номер столбца", "Статичный текст"),
                                     show="headings")

                for i in table_data:
                    table.insert("", tk.END, values=i)

                table.column("Файл", width=105, anchor=tk.CENTER)
                table.column("Порядковый номер таблицы", width=75, anchor=tk.CENTER)
                table.column("Номер строки*Номер столбца", width=130, anchor=tk.CENTER)
                table.column("Статичный текст", width=110, anchor=tk.CENTER)

                table.heading("Файл", text="Файл")
                table.heading("Порядковый номер таблицы", text="№ таблицы")
                table.heading("Номер строки*Номер столбца", text="№ строки и столбца")
                table.heading("Статичный текст", text="Статичный текст")

                scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=table.yview)

                table.configure(yscrollcommand=scrollbar.set)

                table.place(x=10, y=160, height=110)
                scrollbar.place(x=435, y=160, height=110)
            else:
                tk.Label(self.root, text="Не удалось получить данные или файл пуст").place(x=90, y=90)

        def add_template(file, left, right):
            self.clear_window()
            self.templates()
            frame = tk.Frame(self.root)
            frame.place(x=10, y=150, width=100, height=100)

            tk.Label(self.root, text="Путь к файлу").place(x=90, y=90)
            tk.Label(self.root, text='Разделитель').place(x=287, y=90)
            tk.Label(self.root, text='Слева').place(x=280, y=130)
            tk.Label(self.root, text='Справа').place(x=327, y=130)

            new_template = tk.StringVar()
            separator_left = tk.StringVar()
            separator_right = tk.StringVar()
            new = tk.Entry(self.root, textvariable=new_template)
            new.insert(0, file)
            sep_left = tk.Entry(self.root, textvariable=separator_left)
            sep_left.insert(0, left)
            sep_right = tk.Entry(self.root, textvariable=separator_right)
            sep_right.insert(0, right)
            new.place(x=10, y=110, width=250)
            sep_left.place(x=280, y=110, width=40)
            sep_right.place(x=330, y=110, width=40)

            def insert_from_clipboard():
                win32clipboard.OpenClipboard()
                data = win32clipboard.GetClipboardData()
                win32clipboard.CloseClipboard()
                new.insert(0, data.replace('"', '').replace('\\', '/'))

            context_menu = tk.Menu(self.root, tearoff=0)
            context_menu.add_command(label="Вставить из буфера обмена", command=insert_from_clipboard)
            self.root.bind("<Button-3>", lambda event: context_menu.post(event.x_root, event.y_root))

            add_button = tk.Button(self.root, text="Добавить",
                                   command=lambda: find_template(file=new.get(), sep_l=sep_left.get(),
                                                                 sep_r=sep_right.get()))
            add_button.place(x=386, y=107)

        def delete_template(file):
            deleted_file = f'templates/{file}'
            os.remove(deleted_file)
            available_templates_list = find_docx_shablons()
            templates_dropdown = ttk.Combobox(self.root, values=available_templates_list)
            templates_dropdown.set(available_templates_list[0])
            templates_dropdown.place(x=150, y=25)

        def show_template(file):
            def extract_rows_from_database(template_name):
                conn = sqlite3.connect('settings.db')
                c = conn.cursor()

                c.execute(f"SELECT * FROM templates WHERE template LIKE '%{template_name}%'")
                templates_rows = c.fetchall()

                conn.close()

                return templates_rows

            self.clear_window()
            self.templates()

            template_name = file
            rows = extract_rows_from_database(template_name)
            print(rows)
            table = ttk.Treeview(self.root, columns=("Файл", "Старт", "Конец"), show="headings")

            for i in rows:
                table.insert("", tk.END, values=i)

            table.column("Файл", width=120, anchor=tk.CENTER)
            table.column("Старт", width=120, anchor=tk.CENTER)
            table.column("Конец", width=120, anchor=tk.CENTER)

            table.heading("Файл", text="Файл")
            table.heading("Старт", text="Старт")
            table.heading("Конец", text="Конец")

            scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=table.yview)

            table.configure(yscrollcommand=scrollbar.set)

            table.place(x=10, y=100, height=160)
            scrollbar.place(x=375, y=100, height=160)

        self.clear_window()
        self.main_menu()

        templates_label = tk.Label(self.root, text="Шаблоны")
        templates_label.place(x=10, y=5)

        available_templates_label = tk.Label(self.root, text="Доступные шаблоны")
        available_templates_label.place(x=10, y=25)

        available_templates_list = find_docx_shablons()
        if available_templates_list == []:
            available_templates_list.append('None')
        selected_template = tk.StringVar(self.root)
        selected_template.set(available_templates_list[0])
        templates_dropdown = ttk.Combobox(self.root, values=available_templates_list)
        templates_dropdown.set(available_templates_list[0])
        templates_dropdown.place(x=150, y=25)

        # Кнопки
        add_template_button = tk.Button(self.root, text="Добавить шаблон",
                                        command=lambda: add_template(file='', left='', right=''))
        add_template_button.place(x=10, y=55)
        delete_template_button = tk.Button(self.root, text=" Удалить шаблон ",
                                           command=lambda: delete_template(file=templates_dropdown.get()))
        delete_template_button.place(x=300, y=22)
        show_template_button = tk.Button(self.root, text="Показать шаблон",
                                         command=lambda: show_template(file=templates_dropdown.get()))
        show_template_button.place(x=300, y=55)

    def get_data(self):
        def start(queue):
            for data in queue:
                file, template = data
                template = 'templates/' + template

                conn = sqlite3.connect('settings.db')
                c = conn.cursor()

                c.execute(f"SELECT * FROM templates WHERE template LIKE '%{template}%'")
                templates_hits = c.fetchone()

                conn.commit()
                conn.close()

                try:
                    tables_data = extract_tables_from_docx(file_path=file)
                    template_tables_data = extract_tables_from_docx(file_path=template)
                    tables = list(tables_data.keys())
                    template_tables = list(template_tables_data.keys())
                    table_values = []
                    table_variables = []
                    tt = []
                    t = []
                    for key in tables:
                        t.append(tables_data[key])
                        for k in list(tables_data[key].keys()):
                            if tables_data[key][k] != '' and tables_data[key][k] not in table_values:
                                table_values.append(tables_data[key][k])
                    for key in template_tables:
                        tt.append(template_tables_data[key])
                        for k in list(template_tables_data[key].keys()):
                            if template_tables_data[key][k] != '' and template_tables_data[key][
                                k] not in table_variables:
                                table_variables.append(template_tables_data[key][k])
                    pre_table_data = []
                    for h in list(t[0].keys()):
                        if t[0][h] != '' and tt[0][h] != '':
                            pre_table_data.append(get_data_from_string(template=tt[0][h],
                                                                       string=t[0][h],
                                                                       separator_left=templates_hits[1],
                                                                       separator_right=templates_hits[2]))
                    table_data = []
                    for item in pre_table_data:
                        if item != {}:
                            table_data.append(item)

                    def merge_dicts_from_list(dict_list):
                        result_dict = {}
                        for d in dict_list:
                            result_dict.update(d)
                        return result_dict

                    merged_table_data = merge_dicts_from_list(table_data)
                except:
                    merged_table_data = {}

                try:
                    file_paragraphs = read_paragraphs(doc_path=file)
                    template_paragraphs = read_paragraphs(doc_path=template)
                    i = 0
                    pre_data = []

                    while i < len(file_paragraphs):
                        pre_data.append(get_data_from_string(template=template_paragraphs[i], string=file_paragraphs[i],
                                                             separator_left=templates_hits[1],
                                                             separator_right=templates_hits[2]))
                        i += 1
                    data = []
                    abz_result = {}
                    for item in pre_data:
                        if item != {}:
                            data.append(item)
                    for item in data:
                        keys = list(item.keys())
                        for key in keys:
                            if key not in list(abz_result.keys()):
                                abz_result[key] = item[key]
                except:
                    abz_result = {}

                def merge_dicts(dict1, dict2):
                    return {**dict1, **dict2}

                if '.docx' in file:
                    path = file.replace('.docx', '')
                elif '.doc' in file:
                    path = file.replace('.doc', '')
                else:
                    path = 'results'
                result = merge_dicts(dict1=merged_table_data, dict2=abz_result)
                to_csv(data=result, filename=path + '.csv')
                to_excel(data=result, filename=path + '.xlsx')
                to_json(data=result, filename=path + '.json')

        def add_to_queue(new_file: str, template: str, queue: list):
            def delete_data(event):
                for selected_item2 in table.selection():
                    item = table.item(selected_item2)
                    queue.remove(tuple(item["values"]))
                    row_id = table.focus()
                    table.delete(row_id)
                    print(queue)

            if template != '' and new_file != '':
                if '&' in new_file:
                    for f in new_file.split('&'):
                        queue.append((f, template))
                else:
                    queue.append((new_file, template))

            table = ttk.Treeview(self.root, columns=("Файл", "Шаблон"), show="headings")

            for i in queue:
                table.insert("", tk.END, values=i)

            table.column("Файл", width=210, anchor=tk.CENTER)
            table.column("Шаблон", width=210, anchor=tk.CENTER)

            table.heading("Файл", text="Файл")
            table.heading("Шаблон", text="Шаблон")

            scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=table.yview)

            table.configure(yscrollcommand=scrollbar.set)
            table.bind("<<TreeviewSelect>>", delete_data)

            table.place(x=10, y=120, height=150)
            scrollbar.place(x=435, y=120, height=150)

            file.delete(0, 'end')

        self.clear_window()
        self.main_menu()
        tk.Label(self.root, text="Получить данные").place(x=10, y=5)
        tk.Label(self.root, text="Файл").place(x=90, y=30)
        tk.Label(self.root, text="Шаблон").place(x=300, y=30)

        queue = []
        file = tk.Entry(self.root)
        file.place(x=10, y=50, width=200)
        templates = ttk.Combobox(self.root, values=find_docx_shablons())
        templates.place(x=230, y=50, width=200)
        add = tk.Button(self.root, text='Добавить в очередь',
                        command=lambda: add_to_queue(new_file=file.get(), template=templates.get(),
                                                     queue=queue))
        add.place(x=230, y=80)

        start_button = tk.Button(self.root, text='Старт', command=lambda: start(queue=queue))
        start_button.place(x=388, y=80)

        def insert_from_clipboard():
            win32clipboard.OpenClipboard()
            data = win32clipboard.GetClipboardData()
            win32clipboard.CloseClipboard()
            file.insert(0, data.replace('"', '').replace('\\', '/'))

        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Вставить из буфера обмена", command=insert_from_clipboard)
        self.root.bind("<Button-3>", lambda event: context_menu.post(event.x_root, event.y_root))

    def check(self):
        def check_file(file, template, show, empty):
            if file != '' and template != '':

                template = 'templates/' + template

                if show == 'Таблицы':
                    try:
                        tables_data = extract_tables_from_docx(file_path=file)
                        template_tables_data = extract_tables_from_docx(file_path=template)
                        tables = list(tables_data.keys())
                        template_tables = list(template_tables_data.keys())
                        table_values = []
                        table_variables = []
                        tt = []
                        t = []
                        for key in tables:
                            t.append(tables_data[key])
                            for k in list(tables_data[key].keys()):
                                if tables_data[key][k] != '' and tables_data[key][k] not in table_values:
                                    table_values.append(tables_data[key][k])
                        for key in template_tables:
                            tt.append(template_tables_data[key])
                            for k in list(template_tables_data[key].keys()):
                                if template_tables_data[key][k] != '' and template_tables_data[key][
                                    k] not in table_variables:
                                    table_variables.append(template_tables_data[key][k])
                        data = []

                        for h in list(t[0].keys()):
                            if empty == True:
                                data.append((h, tt[0][h], t[0][h]))
                            else:
                                if t[0][h] != '' and tt[0][h] != '':
                                    data.append((h, tt[0][h], t[0][h]))
                    except:
                        print('Файл(-ы) невозможно прочитать или он(они) не содержат таблиц')
                        data = []

                    table = ttk.Treeview(self.root, columns=("№ строки*столбца", "Шаблон", "Файл"), show="headings")

                    for i in data:
                        table.insert("", tk.END, values=i)

                    table.column("№ строки*столбца", width=140, anchor=tk.CENTER)
                    table.column("Шаблон", width=140, anchor=tk.CENTER)
                    table.column("Файл", width=140, anchor=tk.CENTER)

                    table.heading("№ строки*столбца", text="№ строки*столбца")
                    table.heading("Шаблон", text="Шаблон")
                    table.heading("Файл", text="Файл")

                    scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=table.yview)

                    table.configure(yscrollcommand=scrollbar.set)

                    table.place(x=10, y=75, height=200)
                    scrollbar.place(x=435, y=75, height=200)

                else:
                    file_paragraphs = read_paragraphs(doc_path=file)
                    template_paragraphs = read_paragraphs(doc_path=template)

                    if len(file_paragraphs) != len(template_paragraphs):
                        if len(file_paragraphs) > len(template_paragraphs):
                            while len(file_paragraphs) == len(template_paragraphs):
                                template_paragraphs.append('')
                        else:
                            while len(template_paragraphs) > len(file_paragraphs):
                                file_paragraphs.append('')

                    data = []
                    i = 0
                    while i < len(file_paragraphs):
                        if empty != True:
                            if template_paragraphs[i] != '' and file_paragraphs[i] != '':
                                data.append((template_paragraphs[i], file_paragraphs[i]))
                        else:
                            data.append((template_paragraphs[i], file_paragraphs[i]))
                        i += 1

                    table = ttk.Treeview(self.root, columns=("Шаблон", "Файл"), show="headings")

                    for i in data:
                        table.insert("", tk.END, values=i)

                    table.column("Шаблон", width=210, anchor=tk.CENTER)
                    table.column("Файл", width=210, anchor=tk.CENTER)

                    table.heading("Шаблон", text="Шаблон")
                    table.heading("Файл", text="Файл")

                    scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=table.yview)

                    table.configure(yscrollcommand=scrollbar.set)

                    table.place(x=10, y=75, height=200)
                    scrollbar.place(x=435, y=75, height=200)

        self.clear_window()
        self.main_menu()
        tk.Label(self.root, text="Файл").place(x=70, y=5)
        tk.Label(self.root, text="Шаблон").place(x=235, y=5)
        file = tk.Entry(self.root)
        file.place(x=10, y=25, width=160)
        templates = ttk.Combobox(self.root, values=find_docx_shablons())
        templates.place(x=180, y=25, width=160)
        show = tk.StringVar()
        show.set('Таблицы')
        tables_button = tk.Radiobutton(self.root, text='Таблицы', value='Таблицы', variable=show)
        tables_button.place(x=83, y=46)
        abz_button = tk.Radiobutton(self.root, text='Абзацы', value='Абзацы', variable=show)
        abz_button.place(x=170, y=46)
        empty = tk.BooleanVar()
        empty_button = tk.Checkbutton(self.root, text='Показать пустые значения', variable=empty)
        empty_button.place(x=250, y=47)

        check_button = tk.Button(self.root, text='Проверить',
                                 command=lambda: check_file(file=file.get(), template=templates.get(),
                                                            show=show.get(), empty=empty.get()))
        check_button.place(x=350, y=22)
        tk.Label(self.root, text="Отобразить").place(x=5, y=47)

        def insert_from_clipboard():
            win32clipboard.OpenClipboard()
            data = win32clipboard.GetClipboardData()
            win32clipboard.CloseClipboard()
            file.insert(0, data.replace('"', '').replace('\\', '/'))

        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Вставить из буфера обмена", command=insert_from_clipboard)
        self.root.bind("<Button-3>", lambda event: context_menu.post(event.x_root, event.y_root))

    def clear_window(self):
        for widget in self.root.winfo_children():
            widget.destroy()

    def main_menu(self):
        menu = tk.Menu(self.root)
        menu_button = tk.Menu(menu, tearoff=0)

        menu_button.add_command(label="Шаблоны", command=self.templates)
        menu_button.add_command(label="Получить данные", command=self.get_data)
        menu_button.add_command(label="Проверить документ на соответствие шаблону", command=self.check)

        menu.add_cascade(label='Меню', menu=menu_button)
        self.root.config(menu=menu)


app = App()
